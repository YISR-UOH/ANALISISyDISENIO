from docx import *
from docx.shared import *


def crearDocx(lista):
    document = Document()
    png=document.add_picture('imagenes/imagen.png', width=Inches(1))
    document.add_heading('INFORME DE EXPOSICIÓN', 1)
    datos = lista[0]
    info = lista[1]
    trabajadores = lista[2]
    records1 = (
        ('Región / Inspección', 'Año', 'Nº Fiscalización', 'Hoja', 'de'),
        (info['lblRegion'][0] + info['lblIct'][0], info['lblAgno'][0], info['lblFisc'][0], '1', '')
    )

    table1 = document.add_table(rows=0, cols=5)
    table1.style = 'TableGrid'
    for qty, id, id2, id3, id4 in records1:
        row_cells = table1.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = id2
        row_cells[3].text = id3
        row_cells[4].text = id4

    document.add_paragraph('  ')
    document.add_heading('Antecedentes de la empresa y de la fiscalización.', level=1)
    document.add_paragraph('  ')

    records2 = (
        ('Nombre / Razón Social', datos['lblRazonSocial'][0]),
        ('RUT', datos['lblRutEmp'][0]),
        ('Representante Legal', datos['lblNombreRL'][0]),
        ('Rut Rep. Legal', datos['lblRutRL'][0]),
        ('Domicilio Casa Matriz', datos['lblDomL'][0]),
        ('Org. Adm. Ley 16.744',datos['lblOrgAdm16744'][0]),
        ('Total Trabajadores Empresa', trabajadores['lblTotal'][0]),
        ('Fecha Visita', ''),
        ('Hora Inicio', ''),
        ('Fecha de origen', datos['lblFechaOrigen'][0])
    )

    table2 = document.add_table(rows=0, cols=2)
    table2.style = 'TableGrid'
    for qty, id in records2:
        row_cells = table2.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
    document.add_paragraph('  ')
    z=document.add_paragraph('', style='List Number')
    z.add_run('Síntesis del procedimiento y de los medios de investigación utilizados para la constatación de los hechos:').bold = True

    records3 = (
        ('Inicio de fiscalización', 'Con fecha __-__-____, se realiza visita a la dirección indicada en denuncia, específicamente en “empresa”, ubicado en '+datos['lblDomL'][0]),
        ('Materias Fiscalizadas', ''),
        ('Preparación de la visita', ''),
        ('Procedimientos utilizados', ''),
        ('Periodo revisado', ''),
        ('Documentación revisada',''),
        ('Situaciones planteadas por trabajador', ''),
        ('Principio de bilateralidad', ''),
        ('Formularios utilizados', ''),
        ('Visita efectuada en mejor día y hora solicitado por denunciante', ''),
        ('Otro antecedente', '')
    )

    table3 = document.add_table(rows=0, cols=2)
    table3.style = 'TableGrid'
    for qty, id in records3:
        row_cells = table3.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
    document.add_paragraph('  ')

    z=document.add_paragraph('', style='List Number')
    z.add_run('Hechos constatados en relación a las materias fiscalizadas:').bold = True

    document.add_paragraph('  ')

    z=document.add_paragraph('', style='List Number')
    z.add_run('Conclusiones (referidas respecto de la aplicación o no de sanciones):').bold = True


    document.add_page_break()
    return document
